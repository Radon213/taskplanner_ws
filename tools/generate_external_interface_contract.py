#!/usr/bin/env python3
"""Generate the Korean Taskplanner ROS 2 external interface contract DOCX.

The document is built from the checked-in v0.2 visual baseline while all
contract facts and IDL appendices come from the current worktree. Run with the
bundled Codex document Python environment, which provides python-docx.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
import hashlib
from pathlib import Path
import subprocess
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs/Taskplanner_ROS2_External_Interface_Contract_v0.2.0_KO.docx"
OUTPUT = ROOT / "docs/Taskplanner_ROS2_External_Interface_Contract_v0.3.0_KO.docx"

BASELINE = "0.3.0"
SCHEMA_VERSION = "1.1.0"
DATE = date(2026, 8, 13).isoformat()
FONT_BODY = "NanumBarunGothic"
FONT_HEADING = "NanumSquare Neo"
FONT_CODE = "D2Coding ligature"
BLUE = "2E74B5"
NAVY = "1F3A5F"
TEXT = "20262E"
MUTED = "5B636F"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDEBEC"
WHITE = "FFFFFF"
CONTENT_WIDTH_TWIPS = 9360


IDL_FILES = (
    "src/surgical_interop_msgs/msg/GatewayInfo.msg",
    "src/surgical_interop_msgs/msg/ProcedureCatalog.msg",
    "src/surgical_interop_msgs/msg/PhaseCatalogEntry.msg",
    "src/surgical_interop_msgs/msg/InstrumentCatalogEntry.msg",
    "src/surgical_interop_msgs/msg/SurgeryContext.msg",
    "src/surgical_interop_msgs/msg/InstrumentState.msg",
    "src/surgical_interop_msgs/msg/InstrumentStateArray.msg",
    "src/surgical_interop_msgs/msg/RobotState.msg",
    "src/surgical_interop_msgs/msg/RobotStateArray.msg",
    "src/surgical_interop_msgs/msg/RobotEndEffectorState.msg",
    "src/surgical_interop_msgs/msg/RobotEndEffectorStateArray.msg",
    "src/surgical_interop_msgs/msg/ToolPrediction.msg",
    "src/surgical_interop_msgs/msg/ToolPredictionArray.msg",
    "src/surgical_interop_msgs/msg/SpeechRecognitionState.msg",
    "src/surgical_interop_msgs/msg/ClinicalObservation.msg",
    "src/surgical_interop_msgs/msg/ClinicalObservationArray.msg",
    "src/surgical_interop_msgs/msg/SurgeryHealth.msg",
    "src/surgical_interop_msgs/msg/SurgeryEvent.msg",
    "src/surgical_interop_msgs/msg/BedRobotArmState.msg",
    "src/surgical_interop_msgs/msg/BedRobotArmStateArray.msg",
    "src/surgical_interop_msgs/action/ExecuteToolHandover.action",
    "src/surgical_interop_msgs/action/ExecuteRetractionAdjustment.action",
    "src/surgical_interop_msgs/srv/RequestToolChange.srv",
)


TOPICS = (
    ("/surgery/gateway_info", "GatewayInfo", "Gateway heartbeat, schema/interface/catalog and run identity"),
    ("/surgery/catalog", "ProcedureCatalog", "Korean/English procedure, phase, and instrument display metadata"),
    ("/surgery/context", "SurgeryContext", "Active procedure, phase, confidence, execution state, and safety flags"),
    ("/surgery/instruments", "InstrumentStateArray", "Semantic location/state of instrument instances"),
    ("/surgery/robots", "RobotStateArray", "Fresh humanoid and bed-retraction robot execution status"),
    ("/surgery/robot_end_effectors", "RobotEndEffectorStateArray", "Semantic empty/holding/unknown state of robot hands"),
    ("/surgery/tool_predictions", "ToolPredictionArray", "Advisory ranked next-instrument forecast"),
    ("/surgery/speech", "SpeechRecognitionState", "ASR state, sequence, latency, and opt-in finalized free text"),
    ("/surgery/clinical_observations", "ClinicalObservationArray", "Structured model observations; free-text summary is opt-in"),
    ("/surgery/health", "SurgeryHealth", "Source availability, freshness, and stable errors"),
    ("/surgery/events", "SurgeryEvent", "Ordered live state-change facts and explicit outcome"),
)


CATALOG_PHASES = (
    ("1", "P01", "Skin incision", "피부 절개", "P02, P06", "T01, T04, T03"),
    ("2", "P02", "Skin flap elevation", "피부판 거상", "P03, P06", "T03, T05, T06, T02, T04"),
    ("3", "P03", "Strap muscle Exposure", "띠근 노출", "P04, P06", "T05, T02, T04"),
    ("4", "P04", "thyroid Exposure", "갑상선 노출", "P05, P06", "T05, T06, T02, T08, T07"),
    ("5", "P05", "Thyroid vessels ligation", "갑상선 혈관 결찰", "P06", "T02, T08, T09, T07"),
    ("6", "P06", "Bleeding/Hemostasis interrupt", "출혈/지혈 인터럽트", "—", "T10, T07, T04"),
)


CATALOG_TOOLS = (
    ("T01", "#15 Scalpel", "15번 메스", "cutting", "1"),
    ("T02", "Adson forceps", "애드슨 포셉", "grasping", "1"),
    ("T03", "Allis clamp forceps", "알리스 클램프 포셉", "vessel_control", "1"),
    ("T04", "Bovie surgical cautery", "보비 전기소작기", "hemostasis", "1"),
    ("T05", "army navy retractor", "아미-네이비 리트랙터", "exposure", "1"),
    ("T06", "Senn miller retractor", "센 밀러 리트랙터", "exposure", "1"),
    ("T07", "Bipolar cautery", "바이폴라 전기소작기", "hemostasis", "1"),
    ("T08", "Mosquito forceps", "모스키토 포셉", "vessel_control", "1"),
    ("T09", "Harmonics shears", "하모닉 시어", "cutting", "1"),
    ("T10", "Yankeur suction", "양카우어 석션", "suction", "1"),
)


@dataclass(frozen=True)
class TopicDetail:
    title: str
    type_name: str
    role: str
    idle: str
    consumer: str
    example: str


TOPIC_DETAILS = (
    TopicDetail(
        "/surgery/gateway_info",
        "surgical_interop_msgs/msg/GatewayInfo",
        "Gateway 생존, wire/schema 버전, catalog digest, 프로세스 UUID와 활성 run UUID를 제공한다.",
        "heartbeat 유지; procedure_active=false, procedure_run_id는 빈 문자열.",
        "가장 먼저 구독한다. gateway_instance_id가 바뀌면 revision/sequence 캐시를 전부 버린다.",
        """schema_version: \"1.1.0\"
interface_version: \"0.3.0\"
catalog_version: \"sha256:<digest>\"
gateway_instance_id: \"<opaque UUID>\"
procedure_run_id: \"<opaque run UUID>\"
procedure_type: thyroidectomy
procedure_active: true""",
    ),
    TopicDetail(
        "/surgery/catalog",
        "surgical_interop_msgs/msg/ProcedureCatalog",
        "수술·단계·도구의 안정 ID와 영문/한글 표시명, authored 순서, 다음 단계, 예상 도구, alias, category, inventory를 제공한다.",
        "정적 phase/instrument 목록을 계속 제공하되 procedure_active=false.",
        "instrument_id는 procedure_type+catalog_version 범위에서 해석한다. catalog_version 변경 시 label cache를 재생성한다.",
        """procedure_type: thyroidectomy
procedure_display_name: Open Thyroidectomy
procedure_display_name_ko: 갑상선절제술
default_phase_id: P01
phases: [{phase_id: P01, display_name_ko: 피부 절개}]
instruments: [{instrument_id: T04, display_name: Bovie surgical cautery}]""",
    ),
    TopicDetail(
        "/surgery/context",
        "surgical_interop_msgs/msg/SurgeryContext",
        "DT reducer가 수용한 현재 수술 종류·단계·실행 상태·신뢰도·불확실성·safety flag를 제공한다.",
        "procedure_active=false, phase_uncertain=true, evidence_status=UNKNOWN; 나머지 동적 문자열은 비움.",
        "현재 상태이지 미래 단계 예측이 아니다. phase_uncertain과 health를 함께 표시한다.",
        """revision: 1042
procedure_type: thyroidectomy
procedure_active: true
current_phase: P04
phase_confidence: 0.92
phase_uncertain: false
execution_state: running
evidence_status: DT_ACCEPTED
safety_flags: []""",
    ),
    TopicDetail(
        "/surgery/instruments",
        "surgical_interop_msgs/msg/InstrumentStateArray",
        "도구 instance의 의미적 위치, 보유자, lifecycle 상태와 confidence를 전체 snapshot으로 제공한다.",
        "instruments=[]",
        "집도의: holder_role=surgeon + state=handed_over|in_use. "
        "Mayo: location_type=mayo_stand; 회수 여부는 "
        "state=parked_for_reuse|awaiting_retrieval로 구분한다.",
        """revision: 1042
instruments:
  - instrument_id: T04
    instance_id: \"T04#1\"
    location_type: surgeon
    location_id: surgeon
    holder_role: surgeon
    state: in_use
    visible: false
    confidence: 0.96
    evidence_status: DT_ACCEPTED""",
    ),
    TopicDetail(
        "/surgery/robots",
        "surgical_interop_msgs/msg/RobotStateArray",
        "fresh SkillStatus와 controller-owned bed arm status를 read-only robot 상태로 투영한다.",
        "robots=[]",
        "row 부재를 standby로 추론하지 않는다. connection_state=unknown은 source 계약에 신뢰 가능한 연결 bool이 없다는 뜻이다.",
        """revision: 1042
robots:
  - robot_id: humanoid
    robot_type: humanoid
    connection_state: unknown
    execution_state: moving_to_target
    active_command_id: handover-0042
    progress: 0.64
    reason_code: \"\"
    evidence_status: DT_ACCEPTED""",
    ),
    TopicDetail(
        "/surgery/robot_end_effectors",
        "surgical_interop_msgs/msg/RobotEndEffectorStateArray",
        "humanoid right_hand/left_hand의 semantic possession과 보유 도구 ID/instance를 제공한다.",
        "end_effectors=[]",
        "state는 empty, holding, unknown 중 하나다. pose·joint·force·trajectory로 해석하면 안 된다.",
        """procedure_active: true
end_effectors:
  - robot_id: humanoid
    end_effector_id: right_hand
    state: holding
    instrument_id: T04
    instance_id: \"T04#1\"
    confidence: 1.0
    evidence_status: DT_ACCEPTED
  - robot_id: humanoid
    end_effector_id: left_hand
    state: empty""",
    ),
    TopicDetail(
        "/surgery/tool_predictions",
        "surgical_interop_msgs/msg/ToolPredictionArray",
        "현재 구현은 reducer가 수용한 next-tool forecast를 confidence 내림차순 Top-3까지 제공한다.",
        "predictions=[]",
        "advisory UI 정보다. Action Goal이나 handover 권한으로 사용하면 안 된다.",
        """procedure_active: true
predictions:
  - rank: 1
    instrument_id: T09
    instance_id: \"\"
    confidence: 0.87
    stability_sec: 3.4
    source: digital_twin
    evidence_status: DT_ACCEPTED
  - rank: 2
    instrument_id: T04
    instance_id: ""
    confidence: 0.73
    stability_sec: 0.0
    source: digital_twin
    evidence_status: DT_ACCEPTED
  - rank: 3
    instrument_id: T07
    instance_id: ""
    confidence: 0.61
    stability_sec: 0.0
    source: digital_twin
    evidence_status: DT_ACCEPTED""",
    ),
    TopicDetail(
        "/surgery/speech",
        "surgical_interop_msgs/msg/SpeechRecognitionState",
        "ASR availability/connection/state, run-local utterance sequence와 측정 latency를 제공한다. 확정 문장 text는 명시적 free-text opt-in일 때만 제공한다.",
        "available=false, connected=false, state=unavailable, text 비움.",
        "기본값에서는 text가 비어도 status/latency는 유효하다. 문장 수락/로봇 실행 결과가 아니며 latency_available=true일 때만 response_latency_ms를 표시한다.",
        """available: true
connected: true
state: listening
utterance_sequence: 7
text: \"\"
latency_available: true
response_latency_ms: 184.6
latency_basis: api_round_trip
source: taskplanner_asr
evidence_status: GATEWAY_OBSERVED_REDACTED""",
    ),
    TopicDetail(
        "/surgery/clinical_observations",
        "surgical_interop_msgs/msg/ClinicalObservationArray",
        "최신 fresh VLM 결과의 단계·도구·위치·gesture·uncertainty를 제공한다. 자유문장 summary는 명시적 free-text opt-in일 때만 제공한다.",
        "observations=[]",
        "기본값에서 summary는 비어도 구조화 배열은 유효하다. 배열 pair의 길이를 확인하고 model evidence 하나만으로 robot을 command하면 안 된다.",
        """revision: 1043
observations:
  - source: cam4_vlm
    summary: \"\"
    phase_ids: [P04]
    phase_confidences: [0.75]
    observed_tool_ids: [T08]
    observed_location_types: [surgeon]
    observed_location_ids: [surgeon]
    observed_confidences: [0.84]
    uncertainty: 0.22
    evidence_status: MODEL_OBSERVED_REDACTED""",
    ),
    TopicDetail(
        "/surgery/health",
        "surgical_interop_msgs/msg/SurgeryHealth",
        "Gateway가 측정한 source availability, stale, stable error code와 overall required-source health를 제공한다.",
        "실제 source 상태를 계속 발행한다.",
        "healthy=true는 optional source까지 모두 정상이라는 뜻이 아니다. unavailable_sources/stale_sources를 source별로 표시한다.",
        """revision: 1043
healthy: false
state: degraded
unavailable_sources: [cam4]
stale_sources: [vlm_result]
error_codes: [vlm_unhealthy]
evidence_status: GATEWAY_OBSERVED""",
    ),
    TopicDetail(
        "/surgery/events",
        "surgical_interop_msgs/msg/SurgeryEvent",
        "active+fresh run에서 발생한 공개 event fact를 process-local sequence와 명시적 outcome으로 즉시 제공한다.",
        "발행하지 않음; 과거 replay 없음.",
        "DT_ACCEPTED는 event fact의 공개 수용이다. 성공/거절/실패는 state를 확인한다. sequence gap이면 snapshot을 재조회한다.",
        """sequence: 287
schema_version: "1.1.0"
catalog_version: "sha256:<catalog digest>"
gateway_instance_id: "<opaque gateway UUID>"
procedure_run_id: "<opaque run UUID>"
procedure_type: thyroidectomy
event_type: PhaseTransitionRejected
subject_type: procedure
subject_id: P04
phase: P04
state: rejected
correlation_id: \"\"
confidence: 0.3
evidence_status: DT_ACCEPTED""",
    ),
)


def _head() -> str:
    return subprocess.check_output(
        ["git", "rev-parse", "--short=12", "HEAD"], cwd=ROOT, text=True
    ).strip()


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _set_run_font(run, name: str, size_pt: float | None = None) -> None:
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)


def _shade(element, fill: str) -> None:
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(twips / 1440)


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _new_decimal_numbering(doc: DocumentType) -> int:
    """Create one independent decimal list so separate workflows restart at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def _set_table_width(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_TWIPS))
    tbl_w.set(qn("w:type"), "dxa")


def _clear_body(doc: DocumentType) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)
    _set_run_font(run, FONT_BODY, 8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


class Writer:
    def __init__(self, doc: DocumentType) -> None:
        self.doc = doc
        self.table_number = 0
        self.numbering_id: int | None = None

    def heading(self, text: str, level: int = 1, *, page_break: bool = False):
        p = self.doc.add_paragraph(text, style=f"Heading {level}")
        # A paragraph-level page break does not create an empty page when the
        # preceding content already ends exactly at a natural page boundary.
        # An inserted break paragraph can do so (for example before section 6).
        if page_break and self.doc.paragraphs:
            p.paragraph_format.page_break_before = True
        _keep_with_next(p)
        return p

    def body(self, text: str = "", *, bold_prefix: str | None = None):
        p = self.doc.add_paragraph(style="Normal")
        if bold_prefix and text.startswith(bold_prefix):
            lead = p.add_run(bold_prefix)
            lead.bold = True
            _set_run_font(lead, FONT_BODY, 11)
            rest = p.add_run(text[len(bold_prefix):])
            _set_run_font(rest, FONT_BODY, 11)
        else:
            r = p.add_run(text)
            _set_run_font(r, FONT_BODY, 11)
        return p

    def bullet(self, text: str, level: int = 0):
        style = "List Bullet" if level == 0 else "List Bullet 2"
        p = self.doc.add_paragraph(style=style)
        r = p.add_run(text)
        _set_run_font(r, FONT_BODY, 11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        return p

    def number(self, text: str, *, restart: bool = False):
        if restart or self.numbering_id is None:
            self.numbering_id = _new_decimal_numbering(self.doc)
        p = self.doc.add_paragraph(style="List Number")
        _apply_numbering(p, self.numbering_id)
        r = p.add_run(text)
        _set_run_font(r, FONT_BODY, 11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        return p

    def caption(self, text: str):
        self.table_number += 1
        p = self.doc.add_paragraph(f"표 {self.table_number}. {text}", style="Contract Table Caption")
        _keep_with_next(p)
        return p

    def table(
        self,
        caption: str,
        headers: Sequence[str],
        rows: Iterable[Sequence[str]],
        widths: Sequence[int],
    ):
        assert sum(widths) == CONTENT_WIDTH_TWIPS, (caption, sum(widths))
        assert len(headers) == len(widths)
        self.caption(caption)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _set_table_width(table)
        header = table.rows[0]
        _repeat_header(header)
        _cant_split(header)
        for index, text in enumerate(headers):
            cell = header.cells[index]
            _set_cell_width(cell, widths[index])
            _set_cell_margins(cell)
            _shade(cell._tc.get_or_add_tcPr(), PALE_BLUE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            r.bold = True
            _set_run_font(r, FONT_BODY, 9)
            r.font.color.rgb = RGBColor.from_string(NAVY)
        for row_values in rows:
            assert len(row_values) == len(headers), (caption, row_values)
            row = table.add_row()
            _cant_split(row)
            for index, text in enumerate(row_values):
                cell = row.cells[index]
                _set_cell_width(cell, widths[index])
                _set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                r = p.add_run(str(text))
                _set_run_font(r, FONT_BODY, 8.5)
                r.font.color.rgb = RGBColor.from_string(TEXT)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        return table

    def code(self, text: str):
        p = self.doc.add_paragraph(style="Contract Code Block")
        p.paragraph_format.keep_together = False
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.left_indent = Pt(7)
        p.paragraph_format.right_indent = Pt(7)
        _shade(p._p.get_or_add_pPr(), PALE_BLUE)
        r = p.add_run(text.rstrip())
        _set_run_font(r, FONT_CODE, 8)
        r.font.color.rgb = RGBColor.from_string(TEXT)
        return p

    def callout(self, title: str, text: str, *, tone: str = "blue"):
        fill = {"blue": PALE_BLUE, "gold": PALE_GOLD, "red": PALE_RED}[tone]
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        _set_table_width(table)
        _cant_split(table.rows[0])
        cell = table.cell(0, 0)
        _set_cell_width(cell, CONTENT_WIDTH_TWIPS)
        _set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
        _shade(cell._tc.get_or_add_tcPr(), fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        _set_run_font(r, FONT_HEADING, 10.5)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.15
        r2 = p2.add_run(text)
        _set_run_font(r2, FONT_BODY, 9.5)
        r2.font.color.rgb = RGBColor.from_string(TEXT)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        return table


def _configure_styles(doc: DocumentType) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT_HEADING
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_BODY
        style.font.size = Pt(11)
    code = doc.styles["Contract Code Block"]
    code.font.name = FONT_CODE
    code.font.size = Pt(8)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.0


def _set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _configure_section(doc: DocumentType) -> None:
    """Configure a headerless cover section.

    Keeping the cover and body in separate sections avoids a LibreOffice
    pagination defect where ``titlePg`` can be applied to naturally flowed
    continuation pages, dropping their header and top margin.
    """
    section = doc.sections[0]
    _set_section_geometry(section)
    section.different_first_page_header_footer = False
    # LibreOffice renders left/even pages through the explicit even-page
    # relationship once a document contains section-specific headers.
    doc.settings.odd_and_even_pages_header_footer = True
    section.header.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()


def _start_body_section(doc: DocumentType) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_geometry(section)
    section.different_first_page_header_footer = False

    for header in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header.is_linked_to_previous = False
        _populate_body_header(header)
    for footer in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer.is_linked_to_previous = False
        _populate_body_footer(footer)


def _populate_body_header(header) -> None:
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )
    r = p.add_run("TASKPLANNER  |  ROS 2 EXTERNAL INTERFACE CONTRACT")
    r.bold = True
    _set_run_font(r, FONT_HEADING, 8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    r2 = p.add_run(f"\tBASELINE {BASELINE}")
    _set_run_font(r2, FONT_BODY, 8)
    r2.font.color.rgb = RGBColor.from_string(MUTED)


def _populate_body_footer(footer) -> None:
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Taskplanner Integration Baseline  |  ")
    _set_run_font(fr, FONT_BODY, 8.5)
    fr.font.color.rgb = RGBColor.from_string(MUTED)
    _add_page_field(fp)


def _cover(doc: DocumentType, head: str) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tag.add_run("ROS 2 INTEGRATION CONTRACT")
    tr.bold = True
    _set_run_font(tr, FONT_HEADING, 11)
    tr.font.color.rgb = RGBColor.from_string("7A5A00")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(10)
    rr = title.add_run("Taskplanner 외부 연동\nROS 2 인터페이스 계약서")
    rr.bold = True
    _set_run_font(rr, FONT_HEADING, 27)
    rr.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    sr = subtitle.add_run("UI/UX용 공개 상태·영상 토픽 및 로봇 기능 Action/Service 개발 명세")
    _set_run_font(sr, FONT_BODY, 13.5)
    sr.font.color.rgb = RGBColor.from_string("1F4D78")

    for label, value, mono in (
        ("계약 기준", "surgical_interop_msgs 0.3.0 / schema 1.1.0", False),
        ("대상 환경", "ROS 2 Jazzy · Cyclone DDS 검증 기준", False),
        ("소스 기준", f"local main@{head} + 동봉 IDL SHA-256", True),
        ("기준일", DATE, False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        a = p.add_run(f"{label}  ")
        a.bold = True
        _set_run_font(a, FONT_BODY, 10)
        a.font.color.rgb = RGBColor.from_string(MUTED)
        b = p.add_run(value)
        _set_run_font(b, FONT_CODE if mono else FONT_BODY, 10)
        b.font.color.rgb = RGBColor.from_string(TEXT)

    for _ in range(2):
        doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    nr = note.add_run(
        "실제 wire 계약의 최종 원본은 동일 버전의 .msg/.action/.srv 정의입니다. "
        "문서·예시와 불일치하면 IDL이 우선합니다."
    )
    nr.bold = True
    _set_run_font(nr, FONT_BODY, 10)
    nr.font.color.rgb = RGBColor.from_string(NAVY)
    _start_body_section(doc)


def _document_control(w: Writer, head: str) -> None:
    w.heading("문서 통제", 1)
    w.table(
        "문서 식별",
        ("항목", "값"),
        (
            ("문서명", "Taskplanner 외부 연동 ROS 2 인터페이스 계약서"),
            ("Baseline", BASELINE),
            ("Schema / package", f"{SCHEMA_VERSION} / surgical_interop_msgs {BASELINE}"),
            ("소스 기준", f"local main@{head} + 부록 A의 IDL SHA-256"),
            ("작성일", DATE),
            ("상태", "UI/UX 기관 연동용 개발 기준"),
        ),
        (2200, 7160),
    )
    w.callout(
        "계약 우선순위",
        "① 동일 버전의 .msg/.action/.srv 정의 ② 본 문서의 runtime·의미 규칙 "
        "③ 예시 순서로 해석합니다. 예시는 비규범 더미 데이터이며 환자·수술 실데이터가 아닙니다.",
    )
    w.heading("규범 용어", 2)
    w.bullet("MUST(필수): wire 호환성, 정보 경계 또는 안전을 위해 반드시 지켜야 하는 조건.")
    w.bullet("SHOULD(권고): 특별한 근거가 없다면 따라야 하는 상호운용 권고.")
    w.bullet("MAY(선택): wire 형식을 변경하지 않는 구현 선택.")
    w.heading("0.3.0 핵심 변경", 2)
    w.bullet("공개 상태 Gateway 기본값을 true로 변경하고 live·simulation/LLM runtime 공통으로 기동.")
    w.bullet("기존 6개에서 11개 공개 상태/event 토픽으로 확장.")
    w.bullet("Gateway/run/catalog identity, 한·영 카탈로그, 다음 도구 예측, robot 손 상태, ASR latency 추가.")
    w.bullet("시나리오 inactive/stale이면 동적 snapshot과 카메라를 fail-closed empty로 전환.")
    w.bullet("speech·clinical 자유문장은 기본 비공개하고 명시적 개발 opt-in에서만 전달.")
    w.bullet("카메라 native 이름을 stable public alias로 demand-driven relay하고 late frame도 차단.")
    w.heading("문서 구성", 2)
    for index, item in enumerate((
        "목적·범위·소유권",
        "인터페이스 전체 목록",
        "네트워크·QoS·시간 계약",
        "identity·run·idle 계약",
        "공개 토픽 상세",
        "카메라 media 계약",
        "갑상선절제술 표시 카탈로그",
        "로봇 기능 Action/Service",
        "UI 소비 흐름과 코드 예시",
        "검증·인수 기준",
        "안전·보안·정보 경계",
        "부록 A. SHA-256 / 부록 B. 전체 public IDL",
    )):
        w.number(item, restart=index == 0)


def _scope(w: Writer) -> None:
    w.heading("1. 목적·범위·소유권", 1, page_break=True)
    w.body(
        "본 계약은 Taskplanner와 외부 UI/UX·로봇 기관이 동일한 이름, 타입, QoS와 의미를 사용하도록 "
        "공개 ROS 2 경계를 고정한다. UI팀은 공개 상태·영상만 구독하고, robot controller팀은 별도 "
        "Action/Service server를 제공한다. 공개 Gateway는 read-only이며 Taskplanner 내부 DT를 수정하지 않는다."
    )
    w.table(
        "기관 간 역할과 권위",
        ("경계", "Taskplanner", "외부 기관"),
        (
            ("11개 /surgery 상태·event", "sole publisher", "subscriber"),
            ("2개 camera alias", "gated relay publisher", "subscriber"),
            ("Robot Action/Service", "client", "authoritative server"),
            ("Bed arm status source", "subscriber/projector", "authoritative publisher"),
            ("Network/ROS Domain", "deployment 제공", "동일 값 적용·검증"),
        ),
        (2500, 3380, 3480),
    )
    w.callout(
        "쓰기를 금지하는 이유",
        "외부 시스템이 공개 상태 이름에 두 번째 publisher를 만들면 transient-local cache와 권위가 충돌합니다. "
        "상태 토픽은 관찰 전용이며 command는 문서화된 Action/Service만 사용합니다.",
        tone="gold",
    )
    w.heading("1.1 공개 경계에 포함되지 않는 정보", 2)
    for item in (
        "환자 식별자, 수술기록 원문 또는 기타 PHI(자유문장 공개는 기본 OFF)",
        "VLM raw_json, prompt, hidden reasoning, planner rationale, actor hidden state",
        "비공개 event detail 전체(공개 correlation용 command_id/task_id만 예외)",
        "robot pose, joint, force, trajectory, grasp/collision 설정과 private tray slot",
        "자격증명, 내부 network topology, controller free-form diagnostic log",
    ):
        w.bullet(item)
    w.body(
        "0.3.0의 Top-3 도구 예측과 semantic robot-hand possession은 명시적으로 검토된 좁은 예외다. "
        "그 때문에 내부 /twin/world_state 전체가 공개되는 것은 아니다."
    )


def _overview(w: Writer) -> None:
    w.heading("2. 인터페이스 전체 목록", 1, page_break=True)
    w.table(
        "Taskplanner → 외부 UI 공개 상태/event",
        ("토픽", "Type", "요약"),
        TOPICS,
        (2900, 2500, 3960),
    )
    w.table(
        "Taskplanner → 외부 UI camera media",
        ("공개 토픽", "Type", "기본 native source"),
        (
            ("/surgery/images/flir/compressed", "sensor_msgs/msg/CompressedImage", "/synced/flir/color/image_raw/compressed"),
            ("/surgery/images/cam4/compressed", "sensor_msgs/msg/CompressedImage", "/synced/cam_4/color/image_raw/compressed"),
        ),
        (3100, 2700, 3560),
    )
    w.table(
        "Taskplanner → 외부 robot 기능 요청",
        ("Endpoint", "종류·Type", "소유"),
        (
            ("/surgery/tool_handover", "Action · ExecuteToolHandover", "외부 humanoid server"),
            ("/surgery/tool_change/request", "Service · RequestToolChange", "외부 bed-arm controller"),
            ("/surgery/retraction/adjust", "Action · ExecuteRetractionAdjustment", "외부 bed-arm controller"),
        ),
        (3100, 3400, 2860),
    )
    w.callout(
        "Action/Service는 토픽이 아닙니다",
        "Goal/Feedback/Result와 request/response를 동일 이름의 일반 topic처럼 직접 발행하지 않습니다. "
        "외부 서버는 surgical_interop_msgs 0.3.0의 생성된 Action/Service 타입을 설치해야 합니다.",
    )


def _transport(w: Writer) -> None:
    w.heading("3. 네트워크·QoS·시간 계약", 1, page_break=True)
    w.heading("3.1 배포 기준", 2)
    w.table(
        "현재 local main 기본값",
        ("항목", "기본", "비고"),
        (
            ("ROS_DOMAIN_ID", "0", "현재 외부 통합 고정값; 양측이 반드시 같아야 함"),
            ("Discovery", "SUBNET (live mode)", "동일 subnet multicast/discovery 허용 필요"),
            ("RMW", "rmw_cyclonedds_cpp", "양측 일치 및 실제 discovery 검증 필요"),
            ("Cyclone profile", "wired NIC + FragmentSize 1344B", "1500-byte MTU에서 IP fragmentation 회피"),
            ("State Gateway", "PUBLISH_SHARED_STATE=true", "live와 simulation/LLM base runtime 공통"),
            ("Public free text", "PUBLISH_SHARED_FREE_TEXT=false", "speech.text·clinical.summary 기본 suppression"),
            ("Browser bridge", "ENABLE_PUBLIC_ROSBRIDGE=true", "dedicated sidecar; loopback 9092 + wired-LAN proxy"),
            ("Camera alias", "PUBLISH_CAMERA_ALIASES=true", "live wrapper 전용; active/stale gate 적용"),
            ("Snapshot period", "1.0 s", "semantic change와 무관한 heartbeat-style cycle"),
            ("World stale", "3.0 s", "초과 시 inactive/empty fail-closed"),
        ),
        (2600, 2280, 4480),
    )
    w.body(
        "현재 외부 통합 runtime, ASR, public bridge와 Debug sidecar는 Domain 0, SUBNET, Cyclone DDS 및 "
        "검토된 wired profile을 함께 사용한다. replay/shadow는 외부 통합 참가자가 아니며 D71/LOCALHOST로 "
        "격리한다. 배포 담당자는 실제 접속 전에 Domain, RMW, discovery range, interface/subnet, MTU와 "
        "방화벽을 하나의 인수표로 제공해야 한다."
    )
    w.heading("3.2 QoS", 2)
    w.table(
        "공개 delivery profile",
        ("대상", "Reliability", "Durability", "History", "소비 규칙"),
        (
            ("10 snapshot", "Reliable", "Transient Local", "Keep Last 1", "late join retained sample 수신"),
            ("/surgery/events", "Reliable", "Volatile", "Keep Last 50", "연결 이후 live event만"),
            ("camera alias", "Best Effort", "Volatile", "Keep Last 5", "latest frames, demand-gated"),
        ),
        (1800, 1500, 1680, 1380, 3000),
    )
    w.code(
        """# Retained snapshot 수신
ros2 topic echo /surgery/gateway_info \\
  surgical_interop_msgs/msg/GatewayInfo \\
  --qos-reliability reliable --qos-durability transient_local --once

# Camera 수신
ros2 topic echo /surgery/images/flir/compressed \\
  sensor_msgs/msg/CompressedImage \\
  --qos-reliability best_effort --qos-durability volatile --once"""
    )
    w.heading("3.3 시간·신선도", 2)
    w.bullet("Gateway freshness는 untrusted source clock이 아니라 local monotonic receipt time으로 판정한다.")
    w.bullet("메시지 stamp는 발생/수신 시각 표시용이며 freshness 판정과 동일하지 않을 수 있다.")
    w.bullet("revision은 timer cycle이며 semantic state version이나 DB transaction 번호가 아니다.")
    w.bullet("event history depth 50은 DDS queue일 뿐 durable audit log가 아니다.")
    w.doc.add_page_break()
    w.heading("3.4 Native DDS와 ROSBridge", 2)
    w.callout(
        "Native DDS 보안 경계",
        "현재 배포의 DDS subnet은 인증/ACL 경계가 아닙니다. 같은 Domain 참가자는 내부 토픽도 "
        "탐색·구독할 수 있고 공개·내부 토픽 이름에 충돌 또는 위조 샘플을 발행할 수 있습니다. "
        "free-text suppression, 카메라 active gate, 9092 allowlist는 Taskplanner 소유 출력만 제한하며 "
        "다른 DDS 참가자의 트래픽을 필터링하지 않습니다. UI 전용 PC는 DDS에 합류시키지 말고 "
        "9092만 사용하십시오. DDS에는 상호 신뢰된 관리형 제어기만 연결하고 Wi-Fi, Tailscale/VPN, "
        "인터넷으로 라우팅하지 마십시오. 비신뢰 DDS에는 ROS 2/DDS Security identity, governance, "
        "permissions가 필요합니다.",
    )
    w.body(
        "Native ROS 2 consumer는 동일 Domain/discovery/RMW network와 surgical_interop_msgs 0.3.0 설치가 필요하다. "
        "브라우저용 공개 endpoint는 전용 sidecar의 127.0.0.1:9092를 지정 유선 인터페이스에만 proxy한다. "
        "Sidecar는 direct TCP peer가 loopback인지도 검사하여 Tailscale/VPN의 loopback DNAT 우회를 차단한다. "
        "현재 기본 유선 주소가 192.168.1.4이면 ws://192.168.1.4:9092를 사용한다. 이 주소는 DHCP/현장 배포에 따라 "
        "달라질 수 있으므로 배포 담당자가 확정한다. 운영/Debug 9090·9091을 외부 UI가 대신 사용하면 안 된다."
    )
    w.table(
        "Public ROSBridge security profile",
        ("항목", "계약"),
        (
            ("Capabilities", "Subscribe only (subscribe/unsubscribe); incoming fragment rejected"),
            ("Allowlist", "11 state/event + 2 gated camera aliases, exact names"),
            ("Denied", "advertise, publish, service, Action, rosapi"),
            ("Exposure", "direct loopback peer only; designated wired subnet proxy is the sole remote ingress"),
            ("Limits", "8 clients; 64 KiB complete-JSON ingress; camera CBOR/10 Hz/KL1; no fragments; 4 MiB logical egress; 512 MiB sidecar"),
            ("Origin", "defense in depth only; not authentication or the network boundary"),
        ),
        (2500, 6860),
    )


def _identity_idle(w: Writer) -> None:
    w.heading("4. Identity·run·idle 계약", 1, page_break=True)
    w.table(
        "v0.3 identity 필드",
        ("필드", "scope", "consumer 동작"),
        (
            ("schema_version", "projection schema", "지원하지 않는 major/schema이면 fail closed"),
            ("interface_version", "installed IDL package", "0.3.0 타입과 일치 확인"),
            ("catalog_version", "catalog content SHA-256", "변경 시 label/alias cache 재구성"),
            ("gateway_instance_id", "Gateway process", "변경 시 revision/sequence cache 폐기"),
            ("procedure_run_id", "one active run", "변경 시 timeline·dynamic state 초기화"),
            ("procedure_type", "selected catalog", "ID는 이 scope에서만 해석"),
            ("procedure_active", "fresh running predicate", "false면 idle UI와 empty semantics 적용"),
        ),
        (2200, 2600, 4560),
    )
    w.callout(
        "Idle과 연결 끊김은 다릅니다",
        "heartbeat가 fresh하고 procedure_active=false이면 정상 idle입니다. heartbeat 자체가 없거나 stale이면 Gateway unavailable입니다. "
        "마지막 화면을 그대로 남겨 active처럼 보이게 하면 안 됩니다.",
        tone="gold",
    )
    idle_rows = (
        ("gateway_info", "heartbeat 유지; active=false, run_id empty"),
        ("catalog", "phase/tool 정적 목록 유지"),
        ("context", "inactive + uncertain + UNKNOWN"),
        ("instruments / robots", "empty arrays"),
        ("end_effectors / predictions", "empty arrays"),
        ("speech", "unavailable, disconnected, empty text"),
        ("clinical", "empty observations"),
        ("events", "no publication"),
        ("health", "actual source health continues"),
        ("camera", "topic publisher discoverable, no source acquisition/frame"),
    )
    w.table("Inactive/stale output matrix", ("대상", "필수 출력"), idle_rows, (3000, 6360))
    w.body(
        "active boundary에서 이전 run의 speech, VLM, robot/controller cache를 제거한다. active WorldState의 procedure_id가 "
        "loaded catalog와 다르면 dynamic data를 공개하지 않고 health에 procedure_catalog_mismatch를 기록한다."
    )
    w.body(
        "SurgeryEvent는 schema/catalog, gateway instance, procedure run, procedure type을 메시지 자체에 포함한다. "
        "첫 Event가 다음 1 Hz heartbeat보다 먼저 도착해도 자체 run ID로 귀속하며, UI는 sequence 단독이 아니라 "
        "(gateway_instance_id, procedure_run_id)로 묶은 뒤 sequence로 정렬한다."
    )
    w.body(
        "모든 공개 confidence·uncertainty는 유한한 [0,1] 값이다. 잘못된 scalar claim은 UNKNOWN 또는 omit 처리하고, "
        "clinical parallel array는 Gateway에서 길이와 각 row를 검증한다. 소비자도 join 전 길이를 다시 확인한다."
    )
    w.heading("4.1 evidence_status", 2)
    w.table(
        "Evidence authority",
        ("값", "의미", "UI 표시 규칙"),
        (
            ("DT_ACCEPTED", "DT가 operating-state/event fact를 수용", "clinical confirmation 또는 성공으로 오표시 금지"),
            ("MODEL_OBSERVED", "model 관찰/추론", "model evidence badge; robot command 근거 단독 사용 금지"),
            ("CLINICIAN_CONFIRMED", "authorized clinician 확인", "소비 시스템의 별도 safety policy 적용"),
            ("GATEWAY_OBSERVED", "Gateway transport/freshness 측정", "surgical claim이 아닌 health evidence"),
            ("GATEWAY_OBSERVED_REDACTED", "speech item 존재·free text suppression", "sequence/state/latency만 표시; transcript 추론 금지"),
            ("MODEL_OBSERVED_REDACTED", "structured model evidence·summary suppression", "typed ID/confidence만 표시; summary 합성 금지"),
            ("UNKNOWN", "missing/stale/insufficient/cleared", "default 추론 금지"),
        ),
        (1900, 3600, 3860),
    )


def _topic_details(w: Writer) -> None:
    w.heading("5. 공개 토픽 상세 계약", 1, page_break=True)
    w.body(
        "아래 예시는 핵심 필드만 보여주는 더미 값이다. 전체 필드·주석·상수는 부록 B의 IDL 원문을 따른다."
    )
    for index, detail in enumerate(TOPIC_DETAILS, start=1):
        w.heading(f"5.{index} {detail.title}", 2)
        w.table(
            f"{detail.title} 소비 계약",
            ("Type", "역할", "Idle/stale", "Consumer 규칙"),
            ((detail.type_name, detail.role, detail.idle, detail.consumer),),
            (2300, 2500, 2000, 2560),
        )
        w.code(detail.example)
        if detail.title in ("/surgery/speech", "/surgery/clinical_observations"):
            w.callout(
                "자유문장 기본 비공개",
                "PUBLISH_SHARED_FREE_TEXT=false가 기본이다. 이때 speech.text와 clinical summary만 비우고 "
                "상태·sequence·latency 및 구조화 ID/confidence는 유지한다. true는 redaction을 수행하지 않는 "
                "명시적 개발 opt-in이므로 비식별 입력·외부기관 승인·로그 보존 정책을 먼저 확정해야 한다.",
                tone="gold",
            )
        if detail.title == "/surgery/events":
            w.callout(
                "거절 event 해석",
                "evidence_status=DT_ACCEPTED는 '거절 이벤트 사실'을 공개 경계가 수용했다는 뜻입니다. "
                "성공 여부는 state를 읽어야 하며 rejected/failed를 성공 UI로 표시하면 안 됩니다.",
                tone="gold",
            )


def _camera(w: Writer) -> None:
    w.heading("6. 카메라 media 계약", 1, page_break=True)
    w.table(
        "Stable camera aliases",
        ("공개 토픽", "Native source", "QoS", "Gate"),
        (
            ("/surgery/images/flir/compressed", "/synced/flir/color/image_raw/compressed", "BestEffort/Volatile/KL5", "fresh matching active + demand"),
            ("/surgery/images/cam4/compressed", "/synced/cam_4/color/image_raw/compressed", "BestEffort/Volatile/KL5", "fresh matching active + demand"),
        ),
        (2800, 2800, 1900, 1860),
    )
    w.body(
        "Relay는 JPEG를 decode, re-encode, resize, 저장 또는 합성하지 않고 message payload를 그대로 전달한다. "
        "구독자가 있을 때만 native subscription을 만들며, WorldState stop/mismatch/stale이면 즉시 release한다. "
        "frame callback도 gate를 재확인해 timer 사이의 late queued frame을 버린다. Public bridge는 모든 camera "
        "subscription을 CBOR로 강제하여 PNG encoder 경로를 노출하지 않는다."
    )
    w.callout(
        "Privacy gate 우회 방지",
        "Native source와 public topic이 같으면 relay를 생략하지 않고 startup을 실패시킵니다. duplicate public alias와 cross-topic cycle도 거부합니다. "
        "따라서 외부 camera owner가 public 이름에 직접 publish해 idle gate를 우회하면 안 됩니다.",
        tone="red",
    )
    # Keep the code sample and all of its usage rules together; otherwise the
    # final timeout rule can become a one-line orphan before section 7.
    w.heading("6.1 Browser/rosbridge 권고", 2, page_break=True)
    w.code(
        """const image = new ROSLIB.Topic({
  ros,
  name: '/surgery/images/flir/compressed',
  messageType: 'sensor_msgs/msg/CompressedImage',
  compression: 'cbor',
  queue_length: 1,
});

image.subscribe((message) => {
  // Uint8Array 또는 base64 payload를 지원하고 최신 frame만 render한다.
});"""
    )
    w.bullet("rate, resolution, frame_id와 format은 external camera source를 따르며 이 계약이 고정하지 않는다.")
    w.bullet(
        "UI가 hidden/background 상태이면 image subscription을 중단하거나 크게 throttle한다. "
        "frame timeout이면 마지막 JPEG를 유지하지 말고 unavailable/stale로 표시한다."
    )


def _catalog(w: Writer) -> None:
    w.heading("7. 갑상선절제술 표시 카탈로그", 1, page_break=True)
    w.body(
        "현재 기본 bundle은 thyroidectomy이며 ProcedureCatalog가 아래 값을 wire로 발행한다. 표는 UI 개발 참고용이며 "
        "runtime에서는 catalog_version과 실제 topic 내용을 권위로 사용한다."
    )
    w.table(
        "Phase catalog",
        ("순서", "ID", "English", "한국어", "Possible next", "Expected tools"),
        CATALOG_PHASES,
        (720, 720, 1900, 1660, 1760, 2600),
    )
    w.doc.add_page_break()
    w.table(
        "Instrument catalog",
        ("ID", "English", "한국어", "Category", "Inventory"),
        CATALOG_TOOLS,
        (900, 2860, 2300, 2100, 1200),
    )
    w.body(
        "aliases는 ProcedureCatalog의 각 InstrumentCatalogEntry에 함께 발행되므로 UI가 자체 hardcode 목록을 만들 필요가 없다. "
        "inventory_count는 authored procedure inventory이며 live availability count가 아니다. 실제 instance 상태는 /surgery/instruments를 사용한다."
    )


def _robot_endpoints(w: Writer) -> None:
    w.heading("8. 로봇 기능 Action/Service 계약", 1, page_break=True)
    w.callout(
        "공통 안전 규칙",
        "모든 request는 caller-provided command_id를 사용합니다. remote timeout 또는 연결 단절은 성공이 아니며, terminal state를 확인할 수 없으면 "
        "UNKNOWN/REMOTE_STATE_UNKNOWN으로 보존하고 새 ordinary command를 보내지 않습니다.",
        tone="gold",
    )
    w.heading("8.1 /surgery/tool_handover", 2)
    w.table(
        "ExecuteToolHandover Goal",
        ("필드", "Type", "규칙"),
        (
            ("command_id", "string", "caller correlation/idempotency identifier"),
            ("instrument_id", "string", "공유 human-readable 실도구명; private Txx code 금지"),
            ("instrument_instance_id", "string", "instance가 알려졌을 때 지정"),
            ("source_location", "string", "tray, mayo, robot, surgeon"),
            ("target_location", "string", "허용 transition만 server가 수락"),
        ),
        (2300, 1300, 5760),
    )
    w.table(
        "허용 위치 transition",
        ("Source", "Target", "의미"),
        (
            ("tray", "robot", "예측/선정 도구 pick-up 후 stable hold"),
            ("mayo", "robot", "재사용 Mayo 도구 pick-up 후 stable hold"),
            ("tray", "surgeon", "direct pick-up and handover"),
            ("robot", "surgeon", "held tool handover"),
            ("robot", "tray", "unused held tool return"),
            ("mayo", "tray", "used tool retrieval"),
        ),
        (1600, 1600, 6160),
    )
    w.body(
        "Feedback state는 moving_to_source, grasping, moving_to_target, waiting_for_takeover, placing, holding, stopping, "
        "retreating, recovering_to_tray 중 하나다. progress는 [0,1] lifecycle estimate이며 pose나 성공 flag가 아니다."
    )
    w.body(
        "Result final_state는 completed, canceled, failed이며 ROS Action terminal status와 일치해야 한다. Cancel은 즉시 stop이 아니라 "
        "verified compensating recovery다. canceled_source_unchanged 또는 canceled_recovered_to_tray를 확인한 뒤에만 다음 명령이 가능하다."
    )
    w.heading("8.2 /surgery/tool_change/request", 2)
    w.table(
        "RequestToolChange",
        ("구분", "필드", "규칙"),
        (
            ("Request", "command_id", "caller correlation ID"),
            ("Request", "arm_id", "arm_1 또는 arm_2"),
            ("Request", "target_tool_id", "thyroid_retractor 또는 army_navy_retractor"),
            ("Response", "success", "boolean outcome"),
            ("Response", "result", "completed, failed, canceled, protective_stop, unknown"),
            ("Response", "reason_code", "stable machine-readable reason"),
        ),
        (1500, 2400, 5460),
    )
    w.doc.add_page_break()
    w.heading("8.3 /surgery/retraction/adjust", 2)
    w.table(
        "ExecuteRetractionAdjustment 핵심 enum",
        ("필드", "허용값/범위"),
        (
            ("adjustment_mode", "single | multi"),
            ("target_retractor_id", "left_malleable | right_malleable | both_malleable"),
            ("direction_frame", "surgeon_view"),
            ("direction", "up | down | left | right | none"),
            ("axis", "left_right | up_down | none"),
            ("distance_mm", "procedure/controller policy 범위; safety controller 최종 권위"),
            ("Feedback.state", "adjusting | recovering"),
            ("Result.final_state", "completed | canceled | failed"),
        ),
        (3000, 6360),
    )


def _ui_examples(w: Writer) -> None:
    w.heading("9. UI 소비 흐름과 코드 예시", 1, page_break=True)
    w.heading("9.1 권장 startup/reconnect flow", 2)
    steps = (
        "배포 담당자가 제공한 Domain/RMW/discovery/interface를 적용하고 IDL 0.3.0이 resolve되는지 확인한다.",
        "gateway_info를 reliable/transient-local로 구독해 heartbeat와 버전을 검증한다.",
        "catalog를 받아 phase/tool label map을 만들고 catalog_version에 묶는다.",
        "health와 필요한 snapshot을 구독한다.",
        "procedure_active=false이면 idle empty UI를 표시하고 이전 run row를 재사용하지 않는다.",
        "각 Event를 메시지 자체의 gateway_instance_id/procedure_run_id로 묶고, 새 run이면 이전 timeline과 dynamic cache를 초기화한다.",
        "event sequence gap이면 event replay를 기대하지 말고 snapshot을 다시 읽는다.",
        "camera view를 열 때만 best-effort/volatile 구독하고 latest-frame-only로 render한다.",
        "heartbeat가 stale이면 전체 public state를 unavailable로 표시하고 2단계부터 재연결한다.",
    )
    for index, step in enumerate(steps):
        w.number(step, restart=index == 0)
    w.heading("9.2 rclpy snapshot subscriber", 2)
    w.code(
        """import rclpy
from rclpy.node import Node
from rclpy.qos import QoSProfile, ReliabilityPolicy, DurabilityPolicy
from surgical_interop_msgs.msg import GatewayInfo

qos = QoSProfile(depth=1)
qos.reliability = ReliabilityPolicy.RELIABLE
qos.durability = DurabilityPolicy.TRANSIENT_LOCAL

class Probe(Node):
    def __init__(self):
        super().__init__('public_ui_probe')
        self.create_subscription(GatewayInfo, '/surgery/gateway_info', self.on_info, qos)

    def on_info(self, msg):
        print(msg.interface_version, msg.procedure_active, msg.procedure_run_id)

rclpy.init()
rclpy.spin(Probe())"""
    )
    w.doc.add_page_break()
    w.heading("9.3 roslibjs state subscription", 2)
    w.code(
        """const ros = new ROSLIB.Ros({ url: 'ws://192.168.1.4:9092' });

const info = new ROSLIB.Topic({
  ros,
  name: '/surgery/gateway_info',
  messageType: 'surgical_interop_msgs/msg/GatewayInfo',
});

let gatewayInstance = null;
info.subscribe((msg) => {
  if (gatewayInstance !== msg.gateway_instance_id) {
    clearAllPublicCaches();
    gatewayInstance = msg.gateway_instance_id;
  }
  renderConnection(msg.procedure_active ? 'ACTIVE' : 'IDLE');
});"""
    )
    w.callout(
        "Browser transport 전제",
        "192.168.1.4는 현재 기본 wired 주소 예시이며 현장에서 달라질 수 있습니다. 포트는 public 전용 9092입니다. "
        "다른 PC에서 HTTP/TCP뿐 아니라 실제 WebSocket 101 및 topic sample 수신까지 검증해야 합니다.",
        tone="gold",
    )


def _verification(w: Writer) -> None:
    w.heading("10. 검증·인수 기준", 1, page_break=True)
    w.table(
        "최소 인수 체크리스트",
        ("영역", "합격 조건", "증거"),
        (
            ("IDL", "surgical_interop_msgs 0.3.0 build/resolve", "ros2 interface show + SHA-256"),
            ("Gateway idle", "11 topic owner 존재; heartbeat/catalog/health + dynamic empty", "topic echo fixture"),
            ("Gateway active", "run ID 생성, same-cycle revision, reviewed values", "active simulation smoke"),
            ("Run boundary", "이전 speech/VLM/robot cache 재노출 없음", "stop/start regression"),
            ("Procedure mismatch", "inactive fail-closed + health error", "mismatch test"),
            ("Event", "rejected/failed outcome 명시; private detail 미노출", "projection tests"),
            ("Free text", "기본 text/summary empty; opt-in만 원문", "default/opt-in regression"),
            ("Camera idle", "matched subscriber가 있어도 native reader/frame 0", "DDS graph + timeout"),
            ("Camera active", "source reader 1, byte-identical frame, KL1", "payload/hash smoke"),
            ("Camera stop/stale", "source release, late frame drop", "transition/stale smoke"),
            ("Native LAN", "상대 PC에서 Domain discovery와 retained sample 수신", "topic info/echo logs"),
            ("Browser LAN", "read-only allowlist, 101, public sample, non-LAN reject", "network acceptance bundle"),
            ("Robot endpoints", "Goal/request, feedback/result, cancel/timeout failure cases", "command_id별 transcript"),
        ),
        (1800, 4800, 2760),
    )
    w.doc.add_page_break()
    w.heading("10.1 권장 evidence bundle", 2)
    for item in (
        "실행 모드·timestamp·ROS_DOMAIN_ID·RMW·discovery range·interface/IP",
        "ros2 node list, topic list, topic info --verbose, action/service info",
        "idle/active/stop snapshot 원문과 gateway_instance_id/procedure_run_id",
        "카메라 subscriber count와 byte-identical frame evidence",
        "외부 PC native DDS 또는 approved ROSBridge 수신 기록",
        "실패/timeout/reconnect/cancel 결과와 command_id",
        "IDL SHA-256와 source/build revision",
    ):
        w.bullet(item)


def _safety(w: Writer) -> None:
    w.heading("11. 안전·보안·정보 경계", 1, page_break=True)
    w.heading("11.1 임상·정보 안전", 2)
    w.bullet("공개 sample과 문서 예시는 비식별 synthetic/development data만 사용한다.")
    w.bullet("PUBLISH_SHARED_FREE_TEXT=false가 기본이며 speech.text와 clinical.summary를 빈 값으로 발행한다.")
    w.bullet("free-text=true는 de-identification이 아니다. 비식별 입력·외부기관 승인·화면/로그 보존 정책이 있을 때만 사용한다.")
    w.bullet("MODEL_OBSERVED와 DT_ACCEPTED를 CLINICIAN_CONFIRMED로 승격하지 않는다.")
    w.bullet("raw model response, prompt, internal logs와 credentials를 public topic에 포함하지 않는다.")
    w.heading("11.2 Robot safety", 2)
    w.bullet("public state/prediction은 command authorization이 아니다.")
    w.bullet("controller가 Goal acceptance, homing, collision, force, E-stop, protective stop의 최종 권위를 가진다.")
    w.bullet("timeout·disconnect·contradictory Result는 success로 추론하지 않는다.")
    w.bullet("Cancel은 terminal recovery 결과를 확인하기 전 완료로 표시하지 않는다.")
    # Keep the network threat boundary together; the preceding safety bullets
    # can grow as deployment hardening evolves.
    w.heading("11.3 Network", 2, page_break=True)
    w.bullet("native DDS와 WebSocket exposure를 별도 threat boundary로 관리한다.")
    w.bullet("native DDS는 현재 인증/ACL 경계가 아니다. 같은 Domain 참가자는 내부 토픽 구독과 동일 이름 publish가 가능하다.")
    w.bullet("UI 전용 PC는 DDS에 합류시키지 않고 9092만 사용한다. DDS subnet은 상호 신뢰된 관리형 제어기로 제한한다.")
    w.bullet("비신뢰 native DDS가 필요하면 ROS 2/DDS Security identity·governance·permissions를 배포한다.")
    w.bullet("browser bridge는 Subscribe capability(subscribe/unsubscribe)만 등록하고 incoming fragment/unknown op를 거부한다.")
    w.bullet("incoming은 64 KiB 이하의 완성 JSON frame만 허용하고 malformed/incomplete 입력은 buffer clear 후 연결을 닫는다.")
    w.bullet("outgoing fragmentation은 비활성화하며 직렬화된 논리 메시지가 4 MiB를 넘으면 frame을 전혀 보내지 않는다.")
    w.bullet("public bridge는 별도 512 MiB sidecar에서 9092만 사용하며 camera queue와 client 수를 강제 제한한다.")
    w.bullet("LAN proxy는 지정 wired interface/subnet만 허용하고 Wi-Fi/Tailscale/다른 NIC 우회를 검증한다.")
    w.bullet("Sidecar는 direct non-loopback peer를 WebSocket upgrade 전에 거부하여 VPN의 loopback DNAT 우회도 차단한다.")
    w.bullet("방화벽 해제는 상호운용성 검증 절차가 아니며 production 기본값으로 사용하지 않는다.")


def _appendix_hashes(w: Writer, head: str) -> None:
    w.heading("부록 A. 계약 파일 SHA-256", 1, page_break=True)
    w.body(
        f"소스 기준은 local main@{head}이며, 아래 해시는 이 문서에 실제로 포함된 worktree IDL bytes를 식별한다. "
        "commit되지 않은 변경이 있는 경우 commit hash만으로 동일성을 주장하면 안 된다."
    )
    rows = []
    for rel in IDL_FILES:
        path = ROOT / rel
        rows.append((rel.removeprefix("src/surgical_interop_msgs/"), _sha256(path)))
    w.table("Public IDL manifest", ("파일", "SHA-256"), rows, (4200, 5160))
    w.body(
        "추가 runtime 의미 계약: docs/SHARED_SURGICAL_STATE_CONTRACT.md · "
        f"SHA-256 {_sha256(ROOT / 'docs/SHARED_SURGICAL_STATE_CONTRACT.md')}"
    )


def _appendix_idl(w: Writer) -> None:
    w.heading("부록 B. ROS IDL 원문", 1, page_break=True)
    w.body(
        "아래 내용은 문서 생성 시점의 파일을 그대로 삽입했다. 주석도 계약 설명의 일부이지만 wire field/type의 최종 권위는 원본 파일이다."
    )
    for index, rel in enumerate(IDL_FILES, start=1):
        path = ROOT / rel
        short = rel.removeprefix("src/surgical_interop_msgs/")
        w.heading(f"B.{index} {short}", 2)
        w.code(path.read_text(encoding="utf-8"))


def build() -> Path:
    head = _head()
    doc = Document(TEMPLATE)
    _clear_body(doc)
    _configure_styles(doc)
    _configure_section(doc)
    props = doc.core_properties
    props.title = "Taskplanner 외부 연동 ROS 2 인터페이스 계약서"
    props.subject = "UI/UX 공개 상태·영상 토픽 및 로봇 Action/Service 계약"
    props.author = "Taskplanner Integration Team"
    props.keywords = "ROS 2, Taskplanner, surgical_interop_msgs, UI, interface contract"
    props.comments = "Generated from local main worktree and authoritative public IDL files."

    _cover(doc, head)
    writer = Writer(doc)
    _document_control(writer, head)
    _scope(writer)
    _overview(writer)
    _transport(writer)
    _identity_idle(writer)
    _topic_details(writer)
    _camera(writer)
    _catalog(writer)
    _robot_endpoints(writer)
    _ui_examples(writer)
    _verification(writer)
    _safety(writer)
    _appendix_hashes(writer, head)
    _appendix_idl(writer)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
